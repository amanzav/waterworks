"""PDF generation module for cover letters"""

import re
import os
import subprocess
import platform
from pathlib import Path
from typing import Optional
from docx import Document
from docx.shared import Pt
from .utils import sanitize_filename


def get_document_name(company: str, job_title: str) -> str:
    """Get document name for a cover letter
    
    Args:
        company: Company name
        job_title: Job title
        
    Returns:
        Sanitized document name (without extension)
    """
    company_clean = sanitize_filename(company)
    title_clean = sanitize_filename(job_title)
    return f"{company_clean}_{title_clean}"


class PDFBuilder:
    """Build PDF cover letters from text"""
    
    def __init__(self, output_dir: Path, template_path: Optional[Path] = None):
        """Initialize PDF builder
        
        Args:
            output_dir: Directory to save PDFs
            template_path: Optional path to DOCX template
        """
        self.output_dir = output_dir
        self.template_path = template_path
        self.output_dir.mkdir(parents=True, exist_ok=True)
    
    def create_cover_letter(
        self,
        company: str,
        job_title: str,
        cover_text: str,
        signature: Optional[str] = None
    ) -> bool:
        """Create a cover letter PDF
        
        Args:
            company: Company name
            job_title: Job title
            cover_text: Cover letter body text
            signature: Signature line (uses config or "Sincerely" if not provided)
            
        Returns:
            True if successful
        """
        doc_name = get_document_name(company, job_title)
        
        try:
            # Create or load document
            if self.template_path and self.template_path.exists():
                document = Document(str(self.template_path))
            else:
                document = Document()
            
            # Add cover letter text
            # Use default signature if none provided
            signature_text = signature if signature else "Sincerely"
            full_text = f"{cover_text}\n\n{signature_text}"
            
            paragraph = document.add_paragraph()
            run = paragraph.add_run(full_text)
            run.font.size = Pt(11)
            run.font.name = "Garamond"
            
            # Save as DOCX
            docx_path = self.output_dir / f"{doc_name}.docx"
            document.save(str(docx_path))
            
            # Convert to PDF
            self._convert_to_pdf(docx_path)
            
            # Remove DOCX file
            docx_path.unlink()
            
            return True
            
        except Exception as e:
            print(f"      ✗ Error creating PDF: {e}")
            return False
    
    def _convert_to_pdf(self, docx_path: Path) -> bool:
        """Convert DOCX to PDF using platform-appropriate method
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            True if successful
        """
        system = platform.system()
        
        # Try platform-specific conversion methods in order of preference
        if system == "Windows":
            return self._convert_windows(docx_path)
        elif system in ["Linux", "Darwin"]:  # Darwin is macOS
            # Try LibreOffice first, then pypandoc as fallback
            if self._convert_libreoffice(docx_path):
                return True
            return self._convert_pypandoc(docx_path)
        else:
            print(f"      ⚠️  Unsupported platform: {system}")
            return self._fallback_manual(docx_path)
    
    def _convert_windows(self, docx_path: Path) -> bool:
        """Convert DOCX to PDF on Windows using docx2pdf
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            True if successful
        """
        try:
            import pythoncom
            from docx2pdf import convert
            
            # Initialize COM for Windows
            pythoncom.CoInitialize()
            
            try:
                # Convert (docx2pdf automatically uses same name with .pdf extension)
                convert(str(docx_path))
                return True
            finally:
                # Always uninitialize COM
                pythoncom.CoUninitialize()
            
        except ImportError as e:
            print(f"      ⚠️  docx2pdf not available: {e}")
            print(f"      💡 Install with: pip install docx2pdf pywin32")
            return self._fallback_manual(docx_path)
        except Exception as e:
            print(f"      ⚠️  PDF conversion failed: {e}")
            return self._fallback_manual(docx_path)
    
    def _convert_libreoffice(self, docx_path: Path) -> bool:
        """Convert DOCX to PDF using LibreOffice command line
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            True if successful
        """
        # Check if LibreOffice is available
        libreoffice_commands = ['libreoffice', 'soffice']
        
        libreoffice_cmd = None
        for cmd in libreoffice_commands:
            try:
                result = subprocess.run(
                    [cmd, '--version'],
                    capture_output=True,
                    timeout=5
                )
                if result.returncode == 0:
                    libreoffice_cmd = cmd
                    break
            except (FileNotFoundError, subprocess.TimeoutExpired):
                continue
        
        if not libreoffice_cmd:
            return False  # LibreOffice not found, will try fallback
        
        try:
            # Convert using LibreOffice
            # --headless: run without GUI
            # --convert-to pdf: output format
            # --outdir: output directory (same as input)
            subprocess.run(
                [
                    libreoffice_cmd,
                    '--headless',
                    '--convert-to', 'pdf',
                    '--outdir', str(docx_path.parent),
                    str(docx_path)
                ],
                check=True,
                capture_output=True,
                timeout=30
            )
            
            # Check if PDF was created
            pdf_path = docx_path.with_suffix('.pdf')
            if pdf_path.exists():
                return True
            else:
                print(f"      ⚠️  LibreOffice conversion did not create PDF")
                return False
                
        except subprocess.TimeoutExpired:
            print(f"      ⚠️  LibreOffice conversion timed out")
            return False
        except subprocess.CalledProcessError as e:
            print(f"      ⚠️  LibreOffice conversion failed: {e}")
            return False
        except Exception as e:
            print(f"      ⚠️  Unexpected error with LibreOffice: {e}")
            return False
    
    def _convert_pypandoc(self, docx_path: Path) -> bool:
        """Convert DOCX to PDF using pypandoc (requires pandoc)
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            True if successful
        """
        try:
            import pypandoc
            
            pdf_path = docx_path.with_suffix('.pdf')
            
            # Convert using pypandoc
            pypandoc.convert_file(
                str(docx_path),
                'pdf',
                outputfile=str(pdf_path),
                extra_args=['--pdf-engine=xelatex']
            )
            
            if pdf_path.exists():
                return True
            else:
                return False
                
        except ImportError:
            # pypandoc not installed, fall back to manual
            return False
        except Exception as e:
            print(f"      ⚠️  pypandoc conversion failed: {e}")
            return False
    
    def _fallback_manual(self, docx_path: Path) -> bool:
        """Fallback when automatic conversion fails - keep DOCX and inform user
        
        Args:
            docx_path: Path to DOCX file
            
        Returns:
            True (DOCX was created, manual conversion needed)
        """
        system = platform.system()
        
        print(f"      ℹ️  DOCX file saved at: {docx_path}")
        print(f"      💡 Manual conversion needed:")
        
        if system == "Darwin":  # macOS
            print(f"          # Install LibreOffice:")
            print(f"          brew install libreoffice")
            print(f"          # Or convert manually with:")
            print(f"          open '{docx_path}' -a 'Microsoft Word'")
        elif system == "Linux":
            print(f"          # Install LibreOffice:")
            print(f"          sudo apt-get install libreoffice  # Ubuntu/Debian")
            print(f"          sudo dnf install libreoffice      # Fedora")
            print(f"          # Or convert manually:")
            print(f"          libreoffice --headless --convert-to pdf '{docx_path}'")
        else:
            print(f"          Open the file and save as PDF manually")
        
        return True  # DOCX was created successfully
